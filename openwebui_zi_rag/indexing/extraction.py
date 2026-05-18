from __future__ import annotations

import gc
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from pathlib import Path
from typing import Any

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".json",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".html",
    ".css",
    ".log",
    ".csv",
    ".rst",
    ".xml",
    ".yaml",
    ".yml",
}
OFFICE_EXTENSIONS = {
    ".doc",
    ".docx",
    ".pdf",
    ".msg",
    ".odt",
    ".rtf",
    ".xlsx",
    ".xls",
    ".xlsm",
    ".xlsb",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | OFFICE_EXTENSIONS
POINT_RE = re.compile(
    r"^\s*(?:п(?:ункт)?\.?\s*)?((?:\d+(?:\.\d+)*[.)]?|[A-Za-zА-Яа-я]\)))\s+",
    re.IGNORECASE,
)
APPLY_MARKERS = {"v", "∨", "✓", "✔", "+", "x", "х", "да", "yes", "true"}
NOT_APPLY_MARKERS = {"-", "–", "—", "нет", "no", "false", "n/a", "na", "н/п"}
EASYOCR_LANGUAGE_ALIASES = {
    "rus": "ru",
    "russian": "ru",
    "eng": "en",
    "english": "en",
    "ukr": "uk",
    "ukrainian": "uk",
    "deu": "de",
    "ger": "de",
    "fra": "fr",
    "fre": "fr",
    "spa": "es",
    "ita": "it",
}
_EASYOCR_READERS: dict[tuple[tuple[str, ...], bool, str, str], Any] = {}


class ExtractionError(RuntimeError):
    pass


def easyocr_reader_count() -> int:
    return len(_EASYOCR_READERS)


def _torch_cuda_memory(torch_module: Any) -> list[dict[str, Any]]:
    cuda = getattr(torch_module, "cuda", None)
    if cuda is None:
        return []
    try:
        if not cuda.is_available():
            return []
        device_count = int(cuda.device_count() or 0)
    except Exception:
        return []

    devices: list[dict[str, Any]] = []
    for device_id in range(device_count):
        try:
            name = str(cuda.get_device_name(device_id))
        except Exception:
            name = f"cuda:{device_id}"
        try:
            allocated = int(cuda.memory_allocated(device_id))
        except Exception:
            allocated = 0
        try:
            reserved = int(cuda.memory_reserved(device_id))
        except Exception:
            reserved = 0
        devices.append(
            {
                "device": device_id,
                "name": name,
                "allocated_mb": round(allocated / 1024 / 1024, 2),
                "reserved_mb": round(reserved / 1024 / 1024, 2),
            }
        )
    return devices


def clear_ocr_gpu_cache(*, unload_readers: bool = True) -> dict[str, Any]:
    started = time.perf_counter()
    torch_module = sys.modules.get("torch")
    readers_before = len(_EASYOCR_READERS)
    memory_before = _torch_cuda_memory(torch_module) if torch_module is not None else []

    if unload_readers:
        _EASYOCR_READERS.clear()
    gc.collect()

    cuda_available = False
    if torch_module is not None:
        cuda = getattr(torch_module, "cuda", None)
        if cuda is not None:
            try:
                cuda_available = bool(cuda.is_available())
            except Exception:
                cuda_available = False
            if cuda_available:
                try:
                    cuda.empty_cache()
                except Exception:
                    pass
                try:
                    cuda.ipc_collect()
                except Exception:
                    pass
    gc.collect()

    memory_after = _torch_cuda_memory(torch_module) if torch_module is not None else []
    before_reserved = sum(float(item.get("reserved_mb") or 0.0) for item in memory_before)
    after_reserved = sum(float(item.get("reserved_mb") or 0.0) for item in memory_after)
    before_allocated = sum(float(item.get("allocated_mb") or 0.0) for item in memory_before)
    after_allocated = sum(float(item.get("allocated_mb") or 0.0) for item in memory_after)

    return {
        "unload_readers": bool(unload_readers),
        "readers_before": readers_before,
        "readers_after": len(_EASYOCR_READERS),
        "torch_loaded": torch_module is not None,
        "cuda_available": cuda_available,
        "memory_before": memory_before,
        "memory_after": memory_after,
        "freed_reserved_mb": round(max(0.0, before_reserved - after_reserved), 2),
        "freed_allocated_mb": round(max(0.0, before_allocated - after_allocated), 2),
        "elapsed_sec": round(time.perf_counter() - started, 6),
    }


def clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def split_text_blocks(text: str) -> list[str]:
    return [clean_text(part) for part in re.split(r"\r?\n+", text or "") if clean_text(part)]


def point_label(text: str) -> str:
    match = POINT_RE.match(text)
    if not match:
        return ""
    value = match.group(1).rstrip(".)")
    return f"пункт {value}"


def table_marker(value: str) -> str:
    token = clean_text(value).strip().lower()
    if not token:
        return ""
    token = token.replace("ё", "е")
    if token in APPLY_MARKERS:
        return "applies"
    if token in NOT_APPLY_MARKERS:
        return "not_applies"
    return ""


def _is_concise_header(value: str) -> bool:
    text = clean_text(value)
    if not text or table_marker(text) or point_label(text):
        return False
    if len(text) > 120:
        return False
    words = re.findall(r"[0-9A-Za-zА-Яа-яЁё]+", text)
    if len(words) > 12:
        return False
    return True


def _header_for_column(previous_rows: list[list[str]], column_index: int) -> str:
    candidates: list[str] = []
    for row in reversed(previous_rows[-8:]):
        if column_index >= len(row):
            continue
        value = clean_text(row[column_index])
        if not _is_concise_header(value):
            continue
        if value not in candidates:
            candidates.insert(0, value)
        if len(candidates) >= 2:
            break
    return " / ".join(candidates[-2:]).strip()


def _format_plain_table_row(cells: list[str]) -> str:
    return " | ".join(cell or "-" for cell in cells)


def _format_descriptive_cells(cells: list[str], marker_columns: set[int], previous_rows: list[list[str]]) -> list[str]:
    descriptive: list[str] = []
    for index, cell in enumerate(cells):
        value = clean_text(cell)
        if not value or index in marker_columns or table_marker(value):
            continue
        header = _header_for_column(previous_rows, index)
        if index >= 2 and header and header.lower() != value.lower() and not value.startswith(f"{header}:"):
            descriptive.append(f"{header}: {value}")
        else:
            descriptive.append(value)
    return descriptive


def normalize_table_row(cells: list[str], previous_rows: list[list[str]] | None = None) -> str:
    cleaned_cells = [clean_text(cell) for cell in cells]
    previous_rows = previous_rows or []
    marker_items = [
        (index, marker)
        for index, cell in enumerate(cleaned_cells)
        if (marker := table_marker(cell))
    ]
    if not marker_items:
        return _format_plain_table_row(cleaned_cells)

    marker_columns = {index for index, _ in marker_items}
    descriptive = _format_descriptive_cells(cleaned_cells, marker_columns, previous_rows)
    if len(descriptive) < 2:
        return _format_plain_table_row(cleaned_cells)

    applies: list[str] = []
    not_applies: list[str] = []
    for index, marker in marker_items:
        header = _header_for_column(previous_rows, index) or f"колонка {index + 1}"
        if marker == "applies":
            applies.append(header)
        elif marker == "not_applies":
            not_applies.append(header)

    parts = [". ".join(descriptive).strip()]
    if applies:
        parts.append(f"Применимо: {', '.join(applies)}.")
    if not_applies:
        parts.append(f"Не применимо: {', '.join(not_applies)}.")
    return " ".join(part for part in parts if part).strip()


def labeled_text(text: str, locators: list[str]) -> str:
    cleaned = clean_text(text)
    if not cleaned:
        return ""
    labels = [item for item in locators if item]
    point = point_label(cleaned)
    if point and not any(item.startswith("пункт ") for item in labels):
        labels.insert(0, point)
    return f"[{'; '.join(labels)}] {cleaned}" if labels else cleaned


def labeled_blocks(text: str, locators: list[str]) -> list[str]:
    return [
        labeled_text(block, [*locators, f"абз. {index}"])
        for index, block in enumerate(split_text_blocks(text), start=1)
    ]


def is_temporary_office_file(path: str | Path) -> bool:
    name = Path(path).name
    lower = name.lower()
    return name.startswith("~$") or (
        lower.startswith(".~lock.") and lower.endswith("#")
    )


def is_supported_file(path: str | Path) -> bool:
    candidate = Path(path)
    return candidate.suffix.lower() in SUPPORTED_EXTENSIONS and not is_temporary_office_file(candidate)


def extract_text(path: str | Path, config: Any = None) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return "\n\n".join(labeled_blocks(file_path.read_text(encoding="utf-8", errors="ignore"), []))
    if suffix == ".pdf":
        return extract_pdf(file_path, config)
    if suffix == ".docx":
        return extract_docx(file_path)
    if suffix in {".xlsx", ".xls", ".xlsm", ".xlsb"}:
        return extract_excel(file_path)
    if suffix == ".msg":
        return extract_msg(file_path, config)
    if suffix in {".doc", ".odt", ".rtf"}:
        return extract_office(file_path, config)
    raise ExtractionError(f"Unsupported file extension: {suffix}")


def extract_pdf(path: Path, config: Any = None) -> str:
    try:
        import pdfplumber
    except Exception as exc:
        if getattr(config, "enable_ocr", False):
            return extract_pdf_ocr(path, config)
        raise ExtractionError("pdfplumber is required for PDF extraction") from exc

    page_texts: list[str] = []
    missing_pages: list[int] = []
    with pdfplumber.open(str(path)) as pdf:
        for index, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            if not text:
                missing_pages.append(index)
            page_texts.append(text)

    if missing_pages and getattr(config, "enable_ocr", False):
        ocr_texts = extract_pdf_ocr_pages(path, missing_pages, config)
        for index, text in ocr_texts.items():
            page_texts[index] = text
    parts: list[str] = []
    for page_index, text in enumerate(page_texts, start=1):
        parts.extend(labeled_blocks(text, [f"стр. {page_index}"]))
    return "\n\n".join(parts).strip()


def extract_pdf_ocr(path: Path, config: Any = None) -> str:
    pages = extract_pdf_ocr_pages(path, None, config)
    parts: list[str] = []
    for page_index in sorted(pages):
        parts.extend(labeled_blocks(pages[page_index], [f"стр. {page_index + 1}"]))
    return "\n\n".join(parts).strip()


def extract_pdf_ocr_pages(
    path: Path, page_indexes: list[int] | None, config: Any = None
) -> dict[int, str]:
    engine = ocr_engine(config)
    if engine in {"easyocr", "gpu", "cuda"}:
        return extract_pdf_ocr_pages_easyocr(path, page_indexes, config)
    if engine in {"tesseract", "cpu"}:
        return extract_pdf_ocr_pages_tesseract(path, page_indexes, config)
    if engine == "auto":
        try:
            return extract_pdf_ocr_pages_easyocr(path, page_indexes, config)
        except ExtractionError:
            return extract_pdf_ocr_pages_tesseract(path, page_indexes, config)
    raise ExtractionError(f"Unsupported OCR engine: {engine}")


def ocr_engine(config: Any = None) -> str:
    return clean_text(getattr(config, "ocr_engine", "easyocr") or "easyocr").lower()


def pdf_render_scale(config: Any = None) -> float:
    try:
        scale = float(getattr(config, "pdf_render_scale", 2.5) or 2.5)
    except (TypeError, ValueError):
        return 2.5
    if not math.isfinite(scale) or scale <= 0:
        return 2.5
    return scale


def easyocr_languages(languages: str | None) -> list[str]:
    tokens = [
        token.strip().lower()
        for token in re.split(r"[\s,+;/]+", languages or "")
        if token.strip()
    ]
    normalized: list[str] = []
    for token in tokens or ["rus", "eng"]:
        value = EASYOCR_LANGUAGE_ALIASES.get(token, token)
        if value and value not in normalized:
            normalized.append(value)
    return normalized or ["ru", "en"]


def easyocr_model_storage_dir(config: Any = None) -> Path:
    configured = clean_text(getattr(config, "ocr_model_storage_dir", "") or "")
    if configured:
        path = Path(configured)
    else:
        storage_dir = getattr(config, "storage_dir", "./openwebui_zi_rag_storage")
        path = Path(storage_dir) / "easyocr_models"
    path = path.expanduser().resolve()
    path.mkdir(parents=True, exist_ok=True)
    return path


def easyocr_gpu_enabled(config: Any = None) -> bool:
    return bool(getattr(config, "ocr_gpu", True))


def get_easyocr_reader(config: Any = None) -> Any:
    languages = easyocr_languages(getattr(config, "ocr_languages", "rus+eng"))
    gpu_enabled = easyocr_gpu_enabled(config)
    device = clean_text(getattr(config, "ocr_gpu_device", "") or "")
    model_dir = easyocr_model_storage_dir(config)

    if gpu_enabled and device.isdigit():
        os.environ.setdefault("CUDA_VISIBLE_DEVICES", device)

    try:
        import torch
    except Exception as exc:
        raise ExtractionError("torch with CUDA support is required for EasyOCR GPU OCR") from exc

    if gpu_enabled and not torch.cuda.is_available():
        raise ExtractionError("EasyOCR GPU OCR requested, but torch.cuda.is_available() is false")

    try:
        import easyocr
    except Exception as exc:
        raise ExtractionError("easyocr is required for GPU OCR") from exc

    gpu_arg: bool | str = gpu_enabled
    if gpu_enabled and device.startswith("cuda"):
        gpu_arg = device

    key = (tuple(languages), gpu_enabled, device, str(model_dir))
    if key not in _EASYOCR_READERS:
        reader = easyocr.Reader(
            languages,
            gpu=gpu_arg,
            model_storage_directory=str(model_dir),
            download_enabled=True,
            verbose=False,
        )
        reader_device = str(getattr(reader, "device", ""))
        if gpu_enabled and "cpu" in reader_device.lower():
            raise ExtractionError(f"EasyOCR initialized on CPU instead of GPU: {reader_device}")
        _EASYOCR_READERS[key] = reader
    return _EASYOCR_READERS[key]


def extract_pdf_ocr_pages_easyocr(
    path: Path, page_indexes: list[int] | None, config: Any = None
) -> dict[int, str]:
    try:
        import numpy as np
        import pypdfium2 as pdfium
        from PIL import ImageOps
    except Exception as exc:
        raise ExtractionError("numpy, pypdfium2 and Pillow are required for EasyOCR PDF OCR") from exc

    reader = get_easyocr_reader(config)
    render_scale = pdf_render_scale(config)
    output: dict[int, str] = {}
    pdf = pdfium.PdfDocument(str(path))
    try:
        indexes = page_indexes if page_indexes is not None else list(range(len(pdf)))
        for page_index in indexes:
            page = pdf[page_index]
            bitmap = None
            try:
                bitmap = page.render(scale=render_scale)
                image = ImageOps.autocontrast(bitmap.to_pil().convert("RGB"))
                results = reader.readtext(np.array(image), detail=0, paragraph=True)
                text = "\n".join(clean_text(item) for item in results if clean_text(item)).strip()
                if text:
                    output[page_index] = text
            finally:
                if bitmap is not None and hasattr(bitmap, "close"):
                    bitmap.close()
                if hasattr(page, "close"):
                    page.close()
    finally:
        if hasattr(pdf, "close"):
            pdf.close()
    return output


def extract_pdf_ocr_pages_tesseract(
    path: Path, page_indexes: list[int] | None, config: Any = None
) -> dict[int, str]:
    if shutil.which("tesseract") is None:
        raise ExtractionError("tesseract binary is required for OCR")
    try:
        import pypdfium2 as pdfium
        import pytesseract
        from PIL import ImageOps
    except Exception as exc:
        raise ExtractionError("pypdfium2, pytesseract and Pillow are required for OCR") from exc

    languages = getattr(config, "ocr_languages", "rus+eng") or "rus+eng"
    render_scale = pdf_render_scale(config)
    output: dict[int, str] = {}
    pdf = pdfium.PdfDocument(str(path))
    try:
        indexes = page_indexes if page_indexes is not None else list(range(len(pdf)))
        for page_index in indexes:
            page = pdf[page_index]
            bitmap = None
            try:
                bitmap = page.render(scale=render_scale)
                image = ImageOps.autocontrast(bitmap.to_pil().convert("L"))
                text = pytesseract.image_to_string(image, lang=languages).strip()
                if text:
                    output[page_index] = text
            finally:
                if bitmap is not None and hasattr(bitmap, "close"):
                    bitmap.close()
                if hasattr(page, "close"):
                    page.close()
    finally:
        if hasattr(pdf, "close"):
            pdf.close()
    return output


def extract_docx(path: Path) -> str:
    try:
        import docx
    except Exception as exc:
        raise ExtractionError("python-docx is required for DOCX extraction") from exc

    def table_lines(table: Any) -> list[str]:
        lines: list[str] = []
        previous_rows: list[list[str]] = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_parts = [clean_text(paragraph.text) for paragraph in cell.paragraphs]
                for nested_table in cell.tables:
                    cell_parts.extend(table_lines(nested_table))
                cell_text = " ".join(part for part in cell_parts if part)
                cells.append(cell_text)
            if any(cells):
                lines.append(normalize_table_row(cells, previous_rows))
                previous_rows.append([clean_text(cell) for cell in cells])
        return lines

    document = docx.Document(str(path))
    parts = []
    paragraph_no = 0
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        paragraph_no += 1
        parts.append(labeled_text(text, [f"абз. {paragraph_no}"]))
    for table_no, table in enumerate(document.tables, start=1):
        lines = table_lines(table)
        if lines:
            parts.append(f"=== Таблица {table_no} ===")
            for row_no, line in enumerate(lines, start=1):
                parts.append(labeled_text(line, [f"таблица {table_no}", f"строка {row_no}"]))
    return "\n\n".join(parts).strip()


def extract_excel(path: Path) -> str:
    try:
        import pandas as pd
    except Exception as exc:
        raise ExtractionError("pandas/openpyxl/xlrd/pyxlsb are required for Excel extraction") from exc

    def clean(value: Any) -> str:
        if value is None:
            return ""
        try:
            if pd.isna(value):
                return ""
        except Exception:
            pass
        text = str(value).replace("\n", " ").strip()
        text = re.sub(r"\s+", " ", text)
        return text[:120]

    xls = pd.ExcelFile(str(path))
    parts = [f"Excel file: {path.name}"]
    for sheet in xls.sheet_names:
        raw = pd.read_excel(xls, sheet_name=sheet, dtype=object, header=None)
        raw = raw.dropna(axis=0, how="all").dropna(axis=1, how="all")
        parts.append(f"=== Sheet: {sheet} ===")
        if raw.empty:
            parts.append("(empty)")
            continue
        rows = raw.values.tolist()
        previous_rows: list[list[str]] = []
        for row_number, row in enumerate(rows, start=1):
            cells = [clean(value) for value in row]
            if any(cells):
                text = normalize_table_row(cells, previous_rows)
                parts.append(labeled_text(text, [f"лист {sheet}", f"строка {row_number}"]))
                previous_rows.append(cells)
    return "\n\n".join(parts).strip()


def extract_msg(path: Path, config: Any = None, depth: int = 0) -> str:
    try:
        import extract_msg as extract_msg_lib
    except Exception as exc:
        raise ExtractionError("extract-msg is required for MSG extraction") from exc

    def clean(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, bytes):
            value = value.decode("utf-8", errors="ignore")
        return clean_text(str(value))

    def html_to_text(value: Any) -> str:
        html = clean(value)
        if not html:
            return ""
        try:
            from bs4 import BeautifulSoup
        except Exception:
            text = re.sub(r"<(br|p|div|tr|li)\b[^>]*>", "\n", html, flags=re.I)
            text = re.sub(r"<[^>]+>", " ", text)
            return re.sub(r"\s+", " ", text).strip()
        return BeautifulSoup(html, "html.parser").get_text("\n", strip=True)

    def attachment_name(attachment: Any, index: int) -> str:
        for attr in ("longFilename", "shortFilename", "displayName", "name"):
            value = clean(getattr(attachment, attr, ""))
            if value:
                return Path(value).name
        return f"attachment-{index}"

    def saved_attachment_paths(result: Any) -> list[Path]:
        candidates: list[Any]
        if isinstance(result, tuple):
            candidates = [result[-1]]
        elif isinstance(result, list):
            candidates = [item[-1] if isinstance(item, tuple) else item for item in result]
        else:
            candidates = [result]
        paths = [Path(candidate) for candidate in candidates if isinstance(candidate, (str, os.PathLike))]
        return [candidate for candidate in paths if candidate.exists() and candidate.is_file()]

    message = extract_msg_lib.Message(str(path))  # type: ignore[no-untyped-call]
    try:
        parts = [f"MSG file: {path.name}"]
        header_fields = [
            ("Subject", getattr(message, "subject", "")),
            ("From", getattr(message, "sender", "")),
            ("To", getattr(message, "to", "")),
            ("Cc", getattr(message, "cc", "")),
            ("Bcc", getattr(message, "bcc", "")),
            ("Date", getattr(message, "date", "")),
        ]
        for label, value in header_fields:
            text = clean(value)
            if text:
                parts.append(f"{label}: {text}")

        body = str(getattr(message, "body", "") or "").strip() or html_to_text(getattr(message, "htmlBody", ""))
        if body:
            parts.append("=== Body ===")
            parts.extend(labeled_blocks(body, ["тело письма"]))

        attachments = list(getattr(message, "attachments", []) or [])
        if attachments and depth < 2:
            with tempfile.TemporaryDirectory(prefix="zi_rag_msg_") as tmp_dir:
                for index, attachment in enumerate(attachments, start=1):
                    if getattr(attachment, "hidden", False):
                        continue
                    name = attachment_name(attachment, index)
                    parts.append(f"=== Attachment: {name} ===")
                    try:
                        result = attachment.save(customPath=tmp_dir, extractEmbedded=True)
                    except Exception as exc:
                        parts.append(f"Attachment extraction failed: {exc}")
                        continue
                    extracted_any = False
                    for saved_path in saved_attachment_paths(result):
                        if is_supported_file(saved_path):
                            try:
                                if saved_path.suffix.lower() == ".msg":
                                    text = extract_msg(saved_path, config, depth + 1)
                                else:
                                    text = extract_text(saved_path, config)
                            except Exception as exc:
                                parts.append(f"{saved_path.name}: extraction failed: {exc}")
                                continue
                            if text:
                                parts.append(text)
                                extracted_any = True
                    if not extracted_any:
                        parts.append("Attachment type is not supported for text extraction")

        return "\n\n".join(part for part in parts if part).strip()
    finally:
        close = getattr(message, "close", None)
        if callable(close):
            close()


def extract_office(path: Path, config: Any = None) -> str:
    soffice = getattr(config, "soffice_path", "") or shutil.which("soffice")
    if not soffice:
        raise ExtractionError("LibreOffice/soffice is required for this document type")
    with tempfile.TemporaryDirectory(prefix="zi_rag_convert_") as tmp_dir:
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            tmp_dir,
            str(path),
        ]
        result = subprocess.run(command, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            raise ExtractionError(result.stderr.strip() or result.stdout.strip())
        converted = Path(tmp_dir) / f"{path.stem}.docx"
        if not converted.exists():
            candidates = list(Path(tmp_dir).glob("*.docx"))
            if not candidates:
                raise ExtractionError("LibreOffice did not produce a DOCX file")
            converted = candidates[0]
        return extract_docx(converted)


def collect_supported_files(root: str | Path, recursive: bool = True) -> list[Path]:
    candidate = Path(root).expanduser().resolve()
    if candidate.is_file():
        return [candidate] if is_supported_file(candidate) else []
    if not candidate.is_dir():
        return []
    iterator = candidate.rglob("*") if recursive else candidate.glob("*")
    return sorted(path for path in iterator if path.is_file() and is_supported_file(path))
